# -*- coding: utf-8 -*-
"""
Generador del Informe Técnico - Mundiales por Confederaciones
Genera un archivo Word (.docx) profesional
Requiere: pip install python-docx
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── ESTILOS GLOBALES ────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.59)
section.page_height = Cm(27.94)
section.left_margin   = Cm(2.54)
section.right_margin  = Cm(2.54)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)

# Fuente por defecto
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

COLOR_TITULO    = RGBColor(0x1B, 0x3A, 0x6B)   # azul marino
COLOR_SUBTITULO = RGBColor(0x2E, 0x74, 0xB5)   # azul medio
COLOR_ACENTO    = RGBColor(0xF5, 0xA6, 0x23)   # dorado
COLOR_TABLA_CAB = RGBColor(0x1B, 0x3A, 0x6B)   # azul marino (cabecera tabla)

def set_cell_bg(cell, hex_color):
    """Pone color de fondo a una celda."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_titulo(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(16)
        run.font.color.rgb = COLOR_TITULO
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        # línea decorativa debajo
        border_p = doc.add_paragraph()
        border_p.paragraph_format.space_before = Pt(0)
        border_p.paragraph_format.space_after  = Pt(8)
        run2 = border_p.add_run('─' * 80)
        run2.font.color.rgb = COLOR_ACENTO
        run2.font.size = Pt(8)
    elif level == 2:
        run.font.size  = Pt(13)
        run.font.color.rgb = COLOR_SUBTITULO
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        run.font.size  = Pt(11)
        run.font.color.rgb = COLOR_TITULO
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(text, bold_prefix=''):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Cabecera
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], '1B3A6B')
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Filas
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = 'EBF2FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            row.cells[ci].text = str(cell_text)
            set_cell_bg(row.cells[ci], bg)
            for para in row.cells[ci].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(10)
    # Anchos de columna
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    doc.add_paragraph()   # espacio después
    return table

# ══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run('INFORME TÉCNICO')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = COLOR_TITULO

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('Historia Interactiva Basada en Datos')
r2.font.size = Pt(16)
r2.font.color.rgb = COLOR_SUBTITULO

p_sub2 = doc.add_paragraph()
p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_sub2.add_run('Mundiales de Fútbol por Confederaciones (1930 – 2022)')
r3.bold = True
r3.font.size = Pt(14)
r3.font.color.rgb = COLOR_TITULO

doc.add_paragraph()
deco = doc.add_paragraph()
deco.alignment = WD_ALIGN_PARAGRAPH.CENTER
rd = deco.add_run('━' * 50)
rd.font.color.rgb = COLOR_ACENTO

doc.add_paragraph()
meta_lines = [
    ('Autor:', 'Carlos Julio Marín Santana'),
    ('Programa:', 'Maestría en Ciencias de Datos'),
    ('Herramienta:', 'Tableau Cloud'),
    ('Fecha:', 'Junio 2026'),
]
for label, value in meta_lines:
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = p_meta.add_run(f'{label} ')
    rb.bold = True
    rb.font.size = Pt(12)
    rv = p_meta.add_run(value)
    rv.font.size = Pt(12)

doc.add_paragraph()
link_p = doc.add_paragraph()
link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rl = link_p.add_run('Dashboard: https://us-east-1.online.tableau.com/#/site/cmarin-07affe76fa/\nviews/MundialesporConfederaciones/DashboardMundiales')
rl.font.size = Pt(9)
rl.font.color.rgb = COLOR_SUBTITULO

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 1 – DESCRIPCIÓN DEL PROBLEMA
# ══════════════════════════════════════════════════════════════════════════════
add_titulo('1. Descripción del Problema y Solución', level=1)

add_titulo('1.1 Planteamiento del Problema', level=2)
add_body(
    'La historia de los Mundiales de Fútbol abarca 92 años de competencia internacional '
    'con datos distribuidos en múltiples fuentes, lo que dificulta comprender patrones '
    'geopolíticos, la dominancia de confederaciones y la evolución histórica de la '
    'participación global. Existe una brecha entre la abundancia de datos disponibles y '
    'la capacidad de los aficionados, académicos y periodistas para extraer conclusiones '
    'significativas de manera visual e intuitiva.'
)

add_titulo('1.2 Pregunta Central de Investigación', level=2)
add_body(
    '¿Cómo se distribuye el éxito y la participación en los Mundiales de Fútbol entre '
    'las confederaciones continentales a lo largo de 22 ediciones (1930–2022)?'
)

add_titulo('1.3 Cómo lo Resuelve el Proyecto', level=2)
add_body(
    'Se desarrolló una historia interactiva en Tableau Cloud con 8 visualizaciones '
    'integradas en un dashboard que permite al usuario explorar, comparar y descubrir '
    'patrones en 22 ediciones mundialistas, 84 países participantes y 490 registros de '
    'participación. La interactividad mediante filtros por confederación permite '
    'personalizar el análisis según el interés del usuario.'
)

add_body('El dashboard responde las siguientes preguntas analíticas:')
bullets_p1 = [
    '¿Qué confederación ha ganado más Mundiales?',
    '¿Cuáles son los países con mayor número de participaciones?',
    '¿Cómo ha evolucionado la participación global a lo largo del tiempo?',
    '¿Cómo se distribuyen geográficamente los países participantes?',
    '¿Cuál es el medallero histórico por confederación?',
]
for b in bullets_p1:
    add_bullet(b)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 2 – DATOS
# ══════════════════════════════════════════════════════════════════════════════
add_titulo('2. Enlace y Descripción de los Datos', level=1)

add_titulo('2.1 Fuentes de Datos', level=2)
add_table(
    ['Fuente', 'Descripción', 'URL'],
    [
        ['Wikipedia', 'Historial completo Mundiales FIFA 1930–2022', 'https://es.wikipedia.org/wiki/Copa_Mundial_de_Fútbol'],
        ['FIFA.com',  'Datos oficiales de confederaciones',          'https://www.fifa.com'],
        ['ISO 3166',  'Códigos de país para visualización geográfica','https://www.iso.org/iso-3166-country-codes.html'],
    ],
    col_widths=[3, 8, 6]
)

add_titulo('2.2 Dataset Generado — Mundiales_Datos.xlsx', level=2)
add_table(
    ['Hoja', 'Filas', 'Columnas', 'Descripción'],
    [
        ['Participaciones',       '490', '14', 'Registro de cada país en cada Mundial'],
        ['Ediciones',             '22',  '11', 'Datos por torneo (sede, campeón, etc.)'],
        ['Resumen_Confederación', '6',   '8',  'Métricas agregadas por confederación'],
        ['Resumen_País',          '84',  '10', 'Métricas agregadas por país'],
        ['Datos_Mapa',            '84',  '6',  'Datos geográficos para el mapa'],
    ],
    col_widths=[4, 2, 3, 8]
)

add_titulo('2.3 Cobertura Temporal', level=2)
add_body(
    '22 ediciones desde Uruguay 1930 hasta Qatar 2022, excluyendo 1942 y 1946 '
    '(suspendidas por la Segunda Guerra Mundial). Total: 84 países únicos, '
    '6 confederaciones, 490 participaciones registradas.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 3 – PREPROCESAMIENTO
# ══════════════════════════════════════════════════════════════════════════════
add_titulo('3. Descripción del Preprocesamiento', level=1)

add_body(
    'Se desarrolló el script Python crear_tableau_mundiales.py que realizó el proceso '
    'ETL completo (Extracción, Transformación y Carga).'
)

add_titulo('3.1 Extracción', level=2)
for b in [
    'Recopilación manual de datos históricos de Wikipedia para las 22 ediciones.',
    'Mapeo de 84 países a sus confederaciones (UEFA, CONMEBOL, CAF, AFC, CONCACAF, OFC).',
    'Asignación de códigos ISO Alpha-3 para visualización geográfica.',
]:
    add_bullet(b)

add_titulo('3.2 Transformación', level=2)
for b in [
    ('Variables binarias creadas: ', 'Es_Campeón, Es_Subcampeón, Es_Tercero, Es_Cuarto (0/1).'),
    ('Codificación ordinal de fases: ', 'Grupos=1, Octavos=2, Cuartos=3, Semis=4, Final=5, Campeón=6.'),
    ('Resolución de variantes históricas: ', '"Checoslovaquia", "Yugoslavia", "República de Irlanda", "Serbia y Montenegro".'),
    ('Campos derivados generados: ', 'En_Final, En_Podio, En_Top4.'),
]:
    add_bullet(b[1], bold_prefix=b[0])

add_titulo('3.3 Carga', level=2)
for b in [
    'Exportación a Excel (.xlsx) con 5 hojas optimizadas para Tableau.',
    'Encabezados en fila 1 (requerimiento de compatibilidad Tableau).',
    'Tipos de datos verificados para roles geográficos (ISO Code → Código de país).',
]:
    add_bullet(b)

add_titulo('3.4 Herramientas', level=2)
add_table(
    ['Herramienta', 'Versión', 'Uso'],
    [
        ['Python',   '3.x',  'Lenguaje principal del script ETL'],
        ['openpyxl', '3.x',  'Generación del archivo Excel'],
        ['pandas',   '2.x',  'Manipulación y transformación de datos'],
        ['Tableau Cloud', 'Web', 'Visualización interactiva y publicación'],
    ],
    col_widths=[4, 3, 10]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 4 – ESCENARIO DE USO
# ══════════════════════════════════════════════════════════════════════════════
add_titulo('4. Escenario de Uso', level=1)

add_titulo('4.1 Usuarios Objetivo', level=2)
add_table(
    ['Perfil de Usuario', 'Necesidad de Información'],
    [
        ['Académicos y estudiantes', 'Contexto histórico y estadístico para investigación'],
        ['Periodistas deportivos',    'Datos precisos para artículos y reportajes'],
        ['Aficionados al fútbol',     'Exploración libre de la historia mundialista'],
        ['Docentes de estadística',   'Caso de estudio real para enseñanza de visualización'],
    ],
    col_widths=[6, 11]
)

add_titulo('4.2 Escenario Narrativo', level=2)
p_escenario = doc.add_paragraph()
p_escenario.paragraph_format.left_indent  = Cm(1)
p_escenario.paragraph_format.right_indent = Cm(1)
r_esc = p_escenario.add_run(
    '"María es profesora universitaria de estadística deportiva. Quiere enseñar a sus '
    'estudiantes sobre la distribución geopolítica del fútbol mundial. Abre el dashboard '
    'y comienza por el mapa global donde visualiza la concentración de participaciones en '
    'Europa y América del Sur. Hace clic en la burbuja de CONMEBOL y todos los gráficos '
    'se filtran mostrando solo países sudamericanos: Brasil lidera con 22 participaciones '
    'y 5 campeonatos. Luego explora la Evolución Histórica para mostrar cómo UEFA domina '
    'numéricamente pero CONMEBOL tiene mayor tasa de éxito relativa. Finalmente consulta '
    'el Historial de Ediciones para revisar el Mundial 1950 donde Uruguay ganó en Brasil."'
)
r_esc.italic = True
r_esc.font.size = Pt(10.5)

add_titulo('4.3 Tareas Habilitadas', level=2)
for b in [
    ('Comparar ', 'campeonatos entre confederaciones (UEFA vs CONMEBOL).'),
    ('Explorar ', 'distribución geográfica de participaciones en el mapa mundial.'),
    ('Identificar ', 'tendencias históricas de participación (1930–2022).'),
    ('Filtrar ', 'datos interactivamente por confederación.'),
    ('Rankear ', 'países por total de participaciones (Top 20).'),
    ('Consultar ', 'el historial completo de las 22 ediciones.'),
]:
    add_bullet(b[1], bold_prefix=b[0])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 5 – ABSTRACCIÓN DE DATOS
# ══════════════════════════════════════════════════════════════════════════════
add_titulo('5. Descripción de la Abstracción de los Datos', level=1)

add_titulo('5.1 Tipo de Dataset', level=2)
add_body(
    'Tablas relacionadas — datos tabulares planos con múltiples granularidades. '
    'La tabla principal (Participaciones) tiene 490 ítems que representan la unidad '
    'de análisis: cada participación de un país en una edición mundialista.'
)

add_titulo('5.2 Atributos por Tipo', level=2)
add_table(
    ['Atributo', 'Tipo', 'Semántica'],
    [
        ['Año',              'Temporal (cuantitativo ordenado)', 'Edición del Mundial (1930–2022)'],
        ['País',             'Categórico nominal',               'País participante (84 valores únicos)'],
        ['Confederación',    'Categórico nominal',               '6 regiones: UEFA, CONMEBOL, CAF, AFC, CONCACAF, OFC'],
        ['Fase',             'Categórico ordinal',               'Grupos < Octavos < Cuartos < Semis < Final < Campeón'],
        ['Fase_Num',         'Cuantitativo discreto (1–6)',       'Codificación numérica de la fase'],
        ['Es_Campeón',       'Cuantitativo binario (0/1)',        'Indicador de título mundial'],
        ['Es_Subcampeón',    'Cuantitativo binario (0/1)',        'Indicador de subcampeonato'],
        ['Es_Tercero',       'Cuantitativo binario (0/1)',        'Indicador de tercer lugar'],
        ['Participaciones',  'Cuantitativo (conteo)',             'Total de Mundiales jugados por país'],
        ['ISO_Code',         'Categórico nominal (geográfico)',   'Código ISO Alpha-3 para mapa'],
        ['Sede',             'Categórico nominal',               'País anfitrión del torneo'],
        ['N_Equipos',        'Cuantitativo discreto',            'Equipos participantes por edición'],
    ],
    col_widths=[4, 5, 8]
)

add_titulo('5.3 Estructura del Dataset', level=2)
add_table(
    ['Magnitud', 'Valor', 'Descripción'],
    [
        ['Ítems totales',     '490',   'Participaciones país-año únicas'],
        ['Ediciones',         '22',    'Torneos desde 1930 hasta 2022'],
        ['Países',            '84',    'Naciones participantes únicas'],
        ['Confederaciones',   '6',     'UEFA, CONMEBOL, CAF, AFC, CONCACAF, OFC'],
        ['Rango temporal',    '92 años','1930–2022 (sin 1942 y 1946)'],
    ],
    col_widths=[5, 3, 9]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 6 – ABSTRACCIÓN DE TAREAS
# ══════════════════════════════════════════════════════════════════════════════
add_titulo('6. Descripción de la Abstracción de las Tareas', level=1)

add_body('Siguiendo el framework de Munzner (2014), las tareas se clasifican en niveles:')

add_titulo('6.1 Tareas de Alto Nivel', level=2)
add_table(
    ['Tarea', 'Nivel', 'Descripción'],
    [
        ['Comparar',  'Intermedio', 'Comparar campeonatos entre UEFA y CONMEBOL'],
        ['Resumir',   'Alto',       'Visión general de 92 años de historia mundialista'],
        ['Buscar',    'Intermedio', 'Encontrar desempeño de una confederación específica'],
        ['Descubrir', 'Alto',       'Identificar patrones no conocidos en la evolución histórica'],
        ['Explorar',  'Alto',       'Navegar libremente sin hipótesis previa'],
    ],
    col_widths=[3, 3, 11]
)

add_titulo('6.2 Tareas de Bajo Nivel (Acciones Concretas)', level=2)
add_table(
    ['Acción', 'Visualización', 'Pregunta que responde'],
    [
        ['Filtrar por confederación',  'Todas (acción interactiva)',  '¿Cómo se ven los datos solo para UEFA?'],
        ['Identificar máximo',         'Campeonatos x Conf',          '¿Qué confederación ganó más Mundiales?'],
        ['Comparar valores',           'Medallero x Conf',            '¿Cuántos oros/platas/bronces tiene cada confederación?'],
        ['Localizar geográficamente',  'Mapa Participaciones',        '¿De dónde vienen los países participantes?'],
        ['Ordenar por valor',          'Top Países',                  '¿Cuáles son los 20 países con más participaciones?'],
        ['Detectar tendencia',         'Evolución Histórica',         '¿Cómo cambió la participación por confederación en el tiempo?'],
        ['Recuperar valor exacto',     'KPIs Confederaciones',        '¿Cuántas participaciones totales tiene CAF?'],
        ['Consultar detalle',          'Historial Ediciones',         '¿Quién ganó el Mundial de 1986?'],
    ],
    col_widths=[4.5, 4.5, 8]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 7 – MARCAS Y CANALES
# ══════════════════════════════════════════════════════════════════════════════
add_titulo('7. Justificación de Marcas y Canales', level=1)

add_titulo('7.1 Marco Teórico', level=2)
add_body(
    'Se aplicaron los principios de codificación visual de Munzner (2014) y Cairo (2012). '
    'La efectividad de los canales visuales varía según el tipo de atributo: '
    'para atributos cuantitativos, la posición y longitud son los más efectivos; '
    'para atributos categóricos, el color (matiz) es el canal más discriminativo.'
)

add_titulo('7.2 Tabla Resumen — Marcas y Canales por Vista', level=2)
add_table(
    ['Vista', 'Marca', 'Canal Principal', 'Canal Secundario', 'Justificación'],
    [
        ['Campeonatos x Conf',   'Línea (barra)', 'Longitud',             'Color (matiz)',  'Comparación de magnitudes entre categorías'],
        ['Participaciones x Conf','Línea (barra)', 'Longitud',            'Color (matiz)',  'Consistencia visual con Campeonatos x Conf'],
        ['Medallero x Conf',     'Área apilada',  'Longitud + Posición',  'Color (matiz)',  'Composición parte-todo por tipo de logro'],
        ['Mapa Participaciones', 'Punto',          'Posición geográfica', 'Tamaño + Color', 'Canal más natural para datos espaciales'],
        ['Top Países',           'Línea (barra)', 'Longitud + Posición',  'Color (matiz)',  'Ranking ordenado de magnitudes (Top 20)'],
        ['Evolución Histórica',  'Área apilada',  'Posición X (tiempo)',  'Color (matiz)',  'Tendencia temporal con composición'],
        ['Historial Ediciones',  'Punto (texto)', 'Posición (tabla)',     'Color (matiz)',  'Lookup de valores exactos nominales'],
        ['KPIs Confederaciones', 'Punto (texto)', 'Posición (tabla)',     'Color (matiz)',  'Múltiples métricas en formato compacto'],
    ],
    col_widths=[4, 3, 3.5, 3.5, 4]
)

add_titulo('7.3 Justificación Detallada por Vista', level=2)

vistas_detalle = [
    (
        'Vista 1: Campeonatos x Conf — Barras Horizontales',
        [
            ('Marca:', 'Líneas (barras) — representan cada confederación como ítem.'),
            ('Canal Longitud:', 'El más efectivo para comparar cuantitativos (rango = 0–12).'),
            ('Canal Color (matiz):', 'Cada confederación tiene color único y consistente en todo el dashboard.'),
            ('Decisión de orientación:', 'Barras horizontales sobre verticales para facilitar lectura de etiquetas largas (ej: "CONMEBOL"). Reduce rotación cognitiva.'),
        ]
    ),
    (
        'Vista 3: Medallero x Conf — Barras Apiladas',
        [
            ('Marca:', 'Áreas apiladas.'),
            ('Canal Longitud apilada:', 'Permite ver tanto el total como la composición interna (Oro/Plata/Bronce).'),
            ('Limitación conocida:', 'La comparación de segmentos no anclados al eje es menos precisa. Mitigada con etiquetas numéricas en cada segmento.'),
            ('Decisión:', 'Justificada porque el objetivo es mostrar COMPOSICIÓN (parte-todo) además de magnitud total.'),
        ]
    ),
    (
        'Vista 4: Mapa de Participaciones — Mapa de Burbujas',
        [
            ('Marca:', 'Puntos (burbujas) posicionados geográficamente.'),
            ('Canal Posición:', 'El canal más efectivo y natural para datos geoespaciales (latitud/longitud).'),
            ('Canal Tamaño:', 'Codifica la magnitud de participaciones — perceptualmente proporcional al área.'),
            ('Canal Color:', 'Indica la confederación — permite ver clusters geopolíticos visualmente.'),
            ('Mapa de burbujas vs coroplético:', 'Elegido porque el interés es la MAGNITUD (cuantitativo), no la densidad territorial. Brasil y Argentina son más relevantes que el tamaño geográfico de sus países.'),
        ]
    ),
    (
        'Vista 6: Evolución Histórica — Barras Apiladas Temporales',
        [
            ('Marca:', 'Áreas apiladas en eje temporal.'),
            ('Canal Posición X:', 'Eje temporal — orden cronológico natural (1930–2022).'),
            ('Canal Color:', 'Contribución por confederación en cada edición.'),
            ('Decisión de etiquetas:', 'Se eliminaron las etiquetas numéricas internas para evitar saturación visual — la tendencia es más importante que los valores exactos.'),
        ]
    ),
    (
        'Vistas 7 y 8: Tablas de Texto',
        [
            ('Justificación:', 'Para tareas de LOOKUP (búsqueda de valores precisos), la tabla de texto supera a cualquier gráfico en efectividad.'),
            ('Historial Ediciones:', 'Permite consultar quién ganó cada Mundial de forma precisa.'),
            ('KPIs Confederaciones:', 'Complementa los gráficos con valores numéricos exactos que la visualización gráfica no ofrece.'),
        ]
    ),
]

for vista_title, details in vistas_detalle:
    add_titulo(vista_title, level=3)
    for label, text in details:
        add_bullet(text, bold_prefix=label + ' ')

add_titulo('7.4 Principios de Diseño Aplicados', level=2)
add_table(
    ['Principio', 'Implementación en el Dashboard'],
    [
        ['Consistencia cromática',    'Cada confederación tiene el mismo color en todas las 8 vistas'],
        ['Interactividad',            'Filtro por selección conecta todas las vistas simultáneamente'],
        ['Jerarquía visual',          'Header prominente → gráficos principales → tablas de detalle'],
        ['Reducción de ruido',        'Eliminación de cuadrículas, etiquetas redundantes y decimales innecesarios'],
        ['Efectividad de canales',    'Longitud para cuantitativos, color para categóricos, posición para espaciales'],
        ['Pre-atención visual',       'Color como canal pre-atentivo facilita identificación rápida de confederaciones'],
    ],
    col_widths=[5, 12]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 8 – CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════════
add_titulo('8. Conclusiones y Hallazgos Principales', level=1)

conclusiones = [
    ('UEFA domina en volumen, CONMEBOL en eficiencia: ',
     'UEFA lidera en participaciones (259) y campeonatos absolutos (12), pero CONMEBOL tiene '
     'mayor tasa de éxito relativa: 10 campeonatos con solo 91 participaciones.'),
    ('Brasil es el país más dominante: ',
     '22 participaciones (el máximo histórico) y 5 campeonatos — único en la historia.'),
    ('Crecimiento sostenido de participación: ',
     'Desde 13 países en 1930 hasta 32 desde 1998, reflejando la globalización del fútbol.'),
    ('Brecha de desarrollo persistente: ',
     'CAF (África) y AFC (Asia) tienen participación creciente pero sin títulos mundiales — '
     'evidencia de la concentración del fútbol de élite en Europa y América del Sur.'),
    ('Concentración geográfica visible: ',
     'El mapa de burbujas revela claramente la concentración suramericana y europea, con '
     'participación esporádica en África, Asia y Oceanía.'),
    ('La interactividad agrega valor analítico: ',
     'Los filtros por confederación permiten descubrir patrones que no son evidentes en '
     'visualizaciones estáticas, habilitando análisis comparativos personalizados.'),
]

for label, text in conclusiones:
    add_bullet(text, bold_prefix=label)

# ══════════════════════════════════════════════════════════════════════════════
#  REFERENCIAS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_titulo('Referencias', level=1)
refs = [
    'Munzner, T. (2014). Visualization Analysis and Design. CRC Press.',
    'Cairo, A. (2012). The Functional Art: An Introduction to Information Graphics and Visualization. New Riders.',
    'Wikipedia (2024). Copa Mundial de Fútbol. https://es.wikipedia.org/wiki/Copa_Mundial_de_F%C3%BAtbol',
    'FIFA (2024). Historia de los Mundiales. https://www.fifa.com',
    'ISO 3166 (2024). Country Codes. https://www.iso.org/iso-3166-country-codes.html',
]
for ref in refs:
    p_ref = doc.add_paragraph(style='List Bullet')
    p_ref.add_run(ref).font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
#  GUARDAR
# ══════════════════════════════════════════════════════════════════════════════
output_path = r'F:\LUEGO BORRAR\PROYECTO FINAL\Informe_Tecnico_Mundiales_Confederaciones.docx'
doc.save(output_path)
print(f'[OK] Informe generado exitosamente en:')
print(f'     {output_path}')
